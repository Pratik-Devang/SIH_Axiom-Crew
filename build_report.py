from pathlib import Path
import re
import math

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_LEADER
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
ASSETS = ROOT / "report_assets"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)
DOCX_PATH = OUT / "Percorsa_SIH26168_Detailed_Technical_Report.docx"

NAVY = "16324F"
BLUE = "2E74B5"
TEAL = "14866D"
ORANGE = "ED7D31"
INK = "1E293B"
MUTED = "5B6777"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F5F1"
LIGHT_ORANGE = "FFF2E8"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D7DEE8"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "8A6400"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name="Aptos", size=None, bold=None, italic=None, color=INK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_rich_run(paragraph, text, base_size=11, color=INK):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, color=NAVY)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, color=color)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def set_keep_together(paragraph):
    paragraph.paragraph_format.keep_together = True


def add_body(doc, text, bold_lead=None, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8):
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, bold=True, color=NAVY)
    add_rich_run(p, text, 11)
    if italic:
        for run in p.runs:
            run.italic = True
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    add_rich_run(p, text, 10.7)
    return p


def add_number(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    add_rich_run(p, text, 10.7)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_rich_run(p, text, {1: 16, 2: 13, 3: 12}.get(level, 11), color=BLUE if level < 3 else NAVY)
    keep_with_next(p)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper() + "  ")
    set_run_font(r, size=10.5, bold=True, color=color)
    add_rich_run(p, text, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_equation(doc, equation, explanation=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3 if explanation else 0)
    r = p.add_run(equation)
    set_run_font(r, name="Consolas", size=10.2, bold=True, color=NAVY)
    if explanation:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(explanation)
        set_run_font(r2, size=9.2, italic=True, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths_dxa, header_fill=NAVY, first_col_bold=False, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r, size=9.2, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 and len(str(text)) < 28 else WD_ALIGN_PARAGRAPH.LEFT
            add_rich_run(p, str(text), font_size)
            if first_col_bold and idx == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def page_break(doc):
    doc.add_page_break()


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 12, 6),
        3: (12, NAVY, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        style = styles[name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
        style.font.size = Pt(10.7)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("PERCORSA  |  SIH26168 TECHNICAL REPORT")
    set_run_font(r, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(5.7)
    right.width = Inches(0.8)
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run("AI-ML Intelligent Dead Reckoning for GNSS-Denied Navigation")
    set_run_font(lr, size=8, color=MUTED)
    add_page_number(right.paragraphs[0])
    for c in table.rows[0].cells:
        set_cell_margins(c, top=0, bottom=0, start=0, end=0)
    set_table_geometry(table, [8200, 1160], indent=0)
    p = footer.paragraphs[0]
    p._element.getparent().remove(p._element)


def create_architecture_figure(path):
    width, height = 2160, 900
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_paths = [
        Path("C:/Windows/Fonts/aptos.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    font_path = next((p for p in font_paths if p.exists()), None)
    bold_paths = [Path("C:/Windows/Fonts/aptos-bold.ttf"), Path("C:/Windows/Fonts/calibrib.ttf"), Path("C:/Windows/Fonts/arialbd.ttf")]
    bold_path = next((p for p in bold_paths if p.exists()), font_path)
    regular = ImageFont.truetype(str(font_path), 27) if font_path else ImageFont.load_default()
    small = ImageFont.truetype(str(font_path), 24) if font_path else ImageFont.load_default()
    bold = ImageFont.truetype(str(bold_path), 28) if bold_path else regular
    title = ImageFont.truetype(str(bold_path), 34) if bold_path else bold

    def centered_text(box, text, font, fill):
        x1, y1, x2, y2 = box
        bb = draw.multiline_textbbox((0, 0), text, font=font, spacing=5, align="center")
        tw, th = bb[2] - bb[0], bb[3] - bb[1]
        draw.multiline_text(((x1+x2-tw)/2, (y1+y2-th)/2), text, font=font, fill=fill, spacing=5, align="center")

    boxes = [
        ((35, 320, 300, 520), "Phone IMU\naccel + gyro", "#EAF2F8", "#2E74B5"),
        ((365, 320, 660, 520), "Calibration\nsync + frames", "#FFF2E8", "#ED7D31"),
        ((745, 175, 1055, 345), "AI motion model\ndelta-s + variance", "#E8F5F1", "#14866D"),
        ((745, 495, 1055, 665), "INS propagation\npose + biases", "#EAF2F8", "#2E74B5"),
        ((1140, 310, 1435, 530), "Error-State EKF\nstate + covariance", "#FDECEC", "#B33A3A"),
        ((1510, 310, 1795, 530), "Vehicle + map\nconstraints", "#E8F5F1", "#14866D"),
        ((1865, 310, 2125, 530), "Android app\npose + confidence", "#EAF2F8", "#2E74B5"),
        ((745, 720, 1055, 850), "GNSS Trust Manager\ntrusted / denied", "#FFF7DB", "#8A6400"),
        ((1510, 720, 1795, 850), "Offline OSM\nHMM matcher", "#F4F6F9", "#5B6777"),
    ]
    for box, label, fill, edge in boxes:
        draw.rounded_rectangle(box, radius=18, fill=fill, outline=edge, width=5)
        centered_text(box, label, bold, "#16324F")

    def arrow(start, end):
        draw.line([start, end], fill="#536273", width=6)
        x1, y1 = start; x2, y2 = end
        ang = math.atan2(y2-y1, x2-x1)
        size = 18
        pts = [(x2, y2), (x2-size*math.cos(ang-0.55), y2-size*math.sin(ang-0.55)),
               (x2-size*math.cos(ang+0.55), y2-size*math.sin(ang+0.55))]
        draw.polygon(pts, fill="#536273")

    for a, b in [
        ((300, 420), (365, 420)), ((660, 420), (745, 270)), ((660, 420), (745, 580)),
        ((1055, 260), (1140, 380)), ((1055, 580), (1140, 465)), ((1435, 420), (1510, 420)),
        ((1795, 420), (1865, 420)), ((1055, 785), (1230, 530)), ((1650, 720), (1650, 530))
    ]:
        arrow(a, b)
    centered_text((80, 25, 2080, 105), "Percorsa: trusted initialization, denied-mode propagation, confidence-aware recovery", title, "#16324F")
    centered_text((620, 850, 1190, 895), "GNSS is fused only when integrity checks pass", small, "#8A6400")
    centered_text((1400, 850, 1920, 895), "Map constraints are gated by confidence", small, "#5B6777")
    image.save(path)


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    picture_run = p.add_run()
    picture_run.add_picture(str(path), width=Inches(6.45))
    for doc_pr in picture_run._element.xpath(".//wp:docPr"):
        doc_pr.set("title", "Percorsa system architecture")
        doc_pr.set("descr", "Flow diagram showing smartphone IMU preprocessing, parallel AI motion estimation and INS propagation, Error-State EKF fusion, GNSS trust management, confidence-gated map constraints, and Android navigation output.")
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(caption)


def cover_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SMART INDIA HACKATHON 2026")
    set_run_font(r, size=12, bold=True, color=ORANGE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PERCORSA")
    set_run_font(r, name="Aptos Display", size=34, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-ML Based Intelligent Dead Reckoning for\nGNSS-Denied Ground-Vehicle Navigation")
    set_run_font(r, name="Aptos Display", size=19, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Detailed Technical Report | Problem Statement SIH26168")
    set_run_font(r, size=11.5, bold=True, color=MUTED)

    add_callout(doc, "Core design commitment", "**Continuous relative navigation during complete GNSS denial**, using smartphone inertial sensing, uncertainty-aware fusion, vehicle constraints and offline road context. **GNSS is not assumed to remain available.**", fill=LIGHT_TEAL, color=TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(68)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by Team Percorsa")
    set_run_font(r, size=11, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("August 2026 | Technical concept and validation plan")
    set_run_font(r, size=10, color=MUTED)
    page_break(doc)


def add_contents(doc):
    add_heading(doc, "Document Map", 1)
    add_body(doc, "This report is organized to move from the **operational problem** to the **estimator design**, then to the **validation evidence required for an SIH-ready prototype**. The headings are native Word headings, so the Navigation pane can be used to jump between sections.", align=WD_ALIGN_PARAGRAPH.LEFT)
    sections = [
        "1. Executive Summary", "2. Problem Definition and Requirements",
        "3. Operational Design Domain and Assumptions", "4. Related Work and Technical Gap",
        "5. End-to-End System Architecture", "6. Sensor Processing and Initialization",
        "7. Inertial Navigation and Error-State EKF", "8. AI Motion Model",
        "9. Dataset, Training and Generalization", "10. GNSS Integrity and Denied-Mode Operation",
        "11. Vehicle Constraints and Map Matching", "12. Android and Edge Implementation",
        "13. Experimental Methodology", "14. Results and Ablation Reporting",
        "15. Feasibility, Risks and Limitations", "16. Privacy, Security and Deployment",
        "17. Implementation Roadmap", "18. Conclusion", "References", "Technical Appendices"
    ]
    for s in sections:
        add_bullet(doc, s)
    page_break(doc)


def build_document():
    doc = Document()
    configure_document(doc)
    architecture_path = ASSETS / "percorsa_architecture.png"
    create_architecture_figure(architecture_path)
    cover_page(doc)
    add_contents(doc)

    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc, "Percorsa is a **smartphone-first intelligent dead-reckoning system** for maintaining a usable ground-vehicle trajectory when satellite positioning is completely unavailable or cannot be trusted. It combines calibrated accelerometer and gyroscope measurements, a compact temporal neural network, strapdown inertial propagation, an **Error-State Extended Kalman Filter (ESKF)**, ground-vehicle motion constraints and confidence-gated offline map matching. The intended output is not only a latitude and longitude estimate. It is a navigation state containing **position, velocity, heading, sensor biases and uncertainty**, together with an explicit operating mode and integrity status.")
    add_body(doc, "The design deliberately separates **GNSS availability** from **GNSS trust**. A receiver can continue producing fixes in an urban canyon or under spoofing even when those fixes are unsafe. Percorsa therefore passes GNSS through an integrity manager before fusion. Once the source becomes degraded, denied or suspicious, the system freezes the last trusted absolute anchor and propagates relative motion using the inertial, learned and road-constrained estimator. When GNSS returns, recovery is gradual and innovation-gated so that a single bad fix cannot snap the route to an incorrect location.")
    add_callout(doc, "Critical limitation", "A self-contained inertial system **cannot cold-start an absolute global position inside a denied zone**. Percorsa requires one of the following: a last trusted GNSS state, a manually provided origin, or another authenticated absolute reference. The project claims **bounded and quantified drift**, not zero drift.", fill=LIGHT_ORANGE, color=GOLD)
    add_heading(doc, "1.1 Proposed contribution", 2)
    for item in [
        "A **complete denied-mode navigation pipeline**, not an isolated speed-prediction model.",
        "A hybrid estimator in which the neural model learns motion cues while the INS and ESKF preserve **physical state consistency and uncertainty propagation**.",
        "A **GNSS Trust Manager** that distinguishes unavailable, degraded and suspicious fixes before fusion.",
        "A confidence-aware road constraint based on **HMM/Viterbi map matching**, not a nearest-road snap.",
        "A realistic smartphone deployment profile at **10 Hz**, with a separate optional external IMU profile that is not presented as a smartphone requirement.",
        "A validation protocol that prevents leakage across trips, devices and geographies and reports both **accuracy and uncertainty calibration**."
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.2 Target success criteria", 2)
    add_table(doc,
              ["Category", "Initial target", "Evidence required"],
              [
                  ["Navigation accuracy", "**Less than 10% endpoint drift** in controlled denials", "Median, p90, p95 and worst-case drift by outage length"],
                  ["Short route interpretation", "Example: **less than 5 m over 50 m**", "Measured endpoint error and trajectory plot"],
                  ["Long route interpretation", "Example: **less than 100 m over 1 km**", "Distance-normalized drift with confidence interval"],
                  ["Update rate", "**10 Hz** smartphone profile", "End-to-end timestamped loop rate and dropped-sample count"],
                  ["Deployment", "No mandatory OBD/CAN or cloud", "On-device Android demonstration with offline maps"],
                  ["Integrity", "Unsafe GNSS and low-confidence map updates rejected", "Detection rate, false alarms and NIS traces"],
              ], [1800, 2600, 4960], first_col_bold=True)

    add_heading(doc, "2. Problem Definition and Requirements", 1)
    add_heading(doc, "2.1 Why GNSS-denied navigation matters", 2)
    add_body(doc, "Road-navigation applications rely on GNSS for absolute position, speed and time. In **tunnels, underground parking, dense urban canyons, forested corridors and deliberate interference zones**, the signal may disappear or become unreliable because of blockage, attenuation, multipath, jamming or spoofing. A conventional application typically freezes, extrapolates with a simplistic constant-speed assumption or snaps to an implausible road. The result can be incorrect routing, lost fleet visibility and reduced safety for emergency or critical-mobility users.")
    add_body(doc, "Low-cost smartphone MEMS sensors are continuously available but have biases and noise. Double integration of acceleration causes position error to grow rapidly, while small gyroscope errors rotate gravity into the horizontal axes and create false acceleration. The engineering problem is therefore to **extract as much reliable relative motion as possible**, constrain that motion using vehicle physics and road topology, and expose uncertainty honestly until an absolute reference becomes trustworthy again.")
    add_heading(doc, "2.2 Functional requirements", 2)
    req_rows = [
        ["FR-01", "Acquire synchronized accelerometer and gyroscope streams and GNSS when available", "Timestamped sensor frames with gap flags"],
        ["FR-02", "Detect trusted, degraded, denied and suspected-spoof GNSS states", "Operating-mode transition log"],
        ["FR-03", "Propagate pose, velocity, heading and biases without GNSS", "Continuous state at target update rate"],
        ["FR-04", "Infer forward displacement and motion state from temporal IMU windows", "Prediction mean, variance and state probabilities"],
        ["FR-05", "Apply vehicle and stop constraints only when their assumptions are valid", "Constraint acceptance/rejection trace"],
        ["FR-06", "Use offline roads as probabilistic context", "Top-k road hypotheses and match confidence"],
        ["FR-07", "Recover smoothly after denial", "No unsafe position jump; uncertainty contracts gradually"],
        ["FR-08", "Present pose, mode, drift estimate and confidence to the user", "Android interface and event log"],
    ]
    add_table(doc, ["ID", "Requirement", "Observable output"], req_rows, [900, 5210, 3250], first_col_bold=True)
    add_heading(doc, "2.3 Non-functional requirements", 2)
    for item in [
        "**Offline-first operation:** core estimation and maps must work without network access.",
        "**Real-time behavior:** the smartphone profile targets a stable 10 Hz estimator loop.",
        "**Resource control:** model size, latency, memory, CPU load, battery drain and thermal behavior must be measured on the target phone.",
        "**Explainable integrity:** every rejected GNSS or map update should have a reason code and innovation statistic.",
        "**Privacy by default:** trip traces remain on device unless the user explicitly consents to export.",
        "**Graceful degradation:** when uncertainty exceeds the safe limit, the system must mark the estimate unreliable instead of presenting false precision."
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Operational Design Domain and Assumptions", 1)
    add_body(doc, "The first release must have a narrow and defensible **Operational Design Domain (ODD)**. Percorsa is intended for road-going ground vehicles whose dominant motion is forward along a locally planar road network. The phone may be mounted in different fixed orientations, but it is assumed to remain rigid relative to the vehicle after alignment. The vehicle is not assumed to have wheel-speed, steering-angle or CAN access.")
    add_table(doc, ["ODD element", "In scope for prototype", "Out of scope or conditional"], [
        ["Vehicle", "Cars and light road vehicles", "Two-wheelers require separate dynamics validation; rail, drone and robot modes are future work"],
        ["Environment", "Tunnels, underpasses, parking structures, urban blockage", "Off-road terrain only after map and constraint redesign"],
        ["Phone", "Android phone with accelerometer and gyroscope", "Loose handheld use is rejected or flagged"],
        ["GNSS", "Trusted before denial and during validated recovery", "No assumption of intermittent fixes inside complete denial"],
        ["Maps", "Pre-downloaded OpenStreetMap road graph", "Map constraints disabled when map quality or match confidence is low"],
        ["Denial duration", "Short and medium controlled intervals evaluated explicitly", "Indefinite unaided navigation is not claimed"],
    ], [1800, 3760, 3800], first_col_bold=True)
    add_heading(doc, "3.1 Observability and claim boundary", 2)
    add_body(doc, "The state is only partially observable during prolonged denial. Without absolute updates, position and yaw uncertainty grow. Non-holonomic constraints help observe lateral and vertical velocity, stop events help correct velocity, and road topology reduces the set of plausible paths, but these mechanisms do not create absolute truth. The report therefore distinguishes **state estimation**, **constraint-induced correction** and **absolute anchoring**.")
    add_callout(doc, "Claim language", "Use **maintains a continuous estimated trajectory with quantified uncertainty**. Avoid claims such as exact positioning, zero drift, universal hardware independence or indefinite navigation without an external reference.", fill=LIGHT_BLUE)

    add_heading(doc, "4. Related Work and Technical Gap", 1)
    add_body(doc, "Classical strapdown INS is interpretable and computationally efficient, but low-cost MEMS errors accumulate rapidly. Pure deep-learning odometry can learn device and motion patterns, but it may fail outside its training distribution and often lacks calibrated uncertainty. Map matching can correct road-level ambiguity, but naive nearest-edge snapping creates catastrophic errors on parallel roads. Percorsa addresses the gap by combining these methods while giving each component a bounded role.")
    add_table(doc, ["Approach", "Strength", "Main weakness", "Percorsa response"], [
        ["Pure smartphone INS", "Physics-based and transparent", "Bias and attitude errors cause rapid drift", "Bias-aware ESKF plus learned motion and constraints"],
        ["AI-only odometry", "Learns vibration and motion patterns", "Distribution shift and inconsistent state", "Physics propagation and uncertainty-gated fusion"],
        ["GNSS/INS loosely coupled", "Strong when GNSS is healthy", "Unsafe if bad GNSS is fused blindly", "Integrity manager before measurement update"],
        ["Nearest-road snap", "Simple and fast", "Fails on parallel or crossing roads", "HMM/Viterbi with top-k hypotheses"],
        ["OBD/CAN dead reckoning", "Reliable wheel-speed information", "Vehicle integration and permissions required", "Phone-only primary profile"],
    ], [1700, 2200, 2600, 2860], first_col_bold=True, font_size=8.8)

    add_heading(doc, "5. End-to-End System Architecture", 1)
    add_figure(doc, architecture_path, "Figure 1. Proposed Percorsa architecture and trust-gated information flow.")
    add_heading(doc, "5.1 Data flow", 2)
    for step in [
        "The Android sensor service records IMU samples using a **monotonic clock**, preserving acquisition timestamps and reporting gaps.",
        "Calibration, resampling, filtering, gravity handling and phone-to-vehicle alignment produce a synchronized vehicle-frame stream.",
        "The **AI motion model** and **INS propagation** process the same window in parallel. Their outputs are treated as correlated, not independent truth sources.",
        "The ESKF maintains the nominal navigation state and a covariance that represents current uncertainty.",
        "The vehicle-constraint engine proposes non-holonomic and stop updates. The map matcher proposes road-relative pseudo-measurements with confidence scores.",
        "The GNSS Trust Manager authorizes absolute updates only when integrity tests pass.",
        "The Android interface receives the pose, covariance, operating mode, drift estimate and diagnostic reasons."
    ]:
        add_number(doc, step)
    add_heading(doc, "5.2 Operating state machine", 2)
    add_table(doc, ["State", "Entry condition", "Estimator behavior", "Exit condition"], [
        ["INITIALIZING", "No trusted absolute anchor or alignment incomplete", "Collect stationary/moving samples; do not claim global pose", "Stable alignment and trusted origin"],
        ["TRUSTED", "GNSS integrity checks pass", "Fuse GNSS with nominal covariance", "Repeated quality or innovation failures"],
        ["DEGRADED", "Quality weak but not fully absent", "Inflate GNSS covariance or reject individual fixes", "Recovery or denial threshold"],
        ["DENIED", "No usable fix", "INS + AI + valid constraints; covariance expands", "Consecutive trusted fixes"],
        ["SUSPECTED_SPOOF", "Position jump or cross-sensor inconsistency", "Reject GNSS and preserve last trusted anchor", "Multi-test recovery sequence"],
        ["UNSAFE", "Estimated uncertainty exceeds application limit", "Continue internal propagation but suppress confident guidance", "New trusted anchor or operator reset"],
    ], [1400, 2500, 3280, 2180], first_col_bold=True, font_size=8.6)

    add_heading(doc, "6. Sensor Processing and Initialization", 1)
    add_heading(doc, "6.1 Sensor inputs", 2)
    add_body(doc, "The minimum input is a **six-channel IMU** containing three-axis accelerometer and three-axis gyroscope measurements. A magnetometer is optional because it can be strongly corrupted by the vehicle body, phone speakers, charging cables and nearby infrastructure. If used, it must be reliability-gated and never treated as an always-valid heading source.")
    add_heading(doc, "6.2 Coordinate frames", 2)
    add_body(doc, "Three frames are defined: **F_p**, the phone sensor frame; **F_v**, the vehicle frame with x forward, y left and z upward; and **F_n**, the local navigation frame, preferably East-North-Up. The phone-to-vehicle rotation R_vp transforms a vector measured in the phone frame into the vehicle frame. This terminology supports mounting adaptation without claiming that phone orientation is irrelevant.")
    add_equation(doc, "a_v = R_vp * a_p     and     omega_v = R_vp * omega_p", "Phone-frame inertial measurements transformed into the vehicle frame.")
    add_heading(doc, "6.3 Calibration and synchronization", 2)
    for item in [
        "Estimate initial gyroscope bias during a detected stationary interval; track residual bias in the ESKF.",
        "Estimate accelerometer offset and scale using multi-orientation calibration where practical; otherwise use factory calibration plus online bias estimation.",
        "Use the sensor event timestamp, not UI or wall-clock arrival time. Resample to the estimator rate using interpolation and **flag gaps instead of silently filling them**.",
        "Reject saturation and implausible spikes. Keep a per-axis quality flag so bad measurements increase process noise or are skipped.",
        "Apply anti-alias filtering before downsampling high-rate sensor events to the 10 Hz prototype loop."
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.4 Mounting alignment", 2)
    add_body(doc, "At startup, gravity provides roll and pitch. Forward-axis alignment can be estimated during a sufficiently straight acceleration segment by comparing horizontal specific force, velocity direction from trusted GNSS and gyroscope yaw rate. The alignment must be revalidated after a stop or when a phone-movement detector observes an abrupt frame change. If the phone is moved during denial, the system should enter **UNSAFE** or request re-alignment rather than silently continuing.")
    add_heading(doc, "6.5 Preprocessing outputs", 2)
    add_table(doc, ["Output", "Units", "Rate", "Quality metadata"], [
        ["Specific force in F_v", "m/s^2", "10 Hz prototype", "saturation, gap, variance"],
        ["Angular rate in F_v", "rad/s", "10 Hz prototype", "bias, saturation, variance"],
        ["Optional magnetic field", "microtesla", "device dependent", "norm and disturbance score"],
        ["Trusted GNSS measurement", "position/velocity", "as available", "accuracy, satellites, C/N0, NIS"],
        ["Motion window", "normalized tensor", "one output per estimator step", "missingness mask and device ID"],
    ], [2200, 1500, 1800, 3860], first_col_bold=True)

    add_heading(doc, "7. Inertial Navigation and Error-State EKF", 1)
    add_heading(doc, "7.1 Nominal and error state", 2)
    add_body(doc, "The nominal state contains position p_n, velocity v_n, orientation quaternion q_np, accelerometer bias b_a and gyroscope bias b_g. A **15-dimensional local error state** represents small perturbations in position, velocity, attitude and both biases. This formulation keeps quaternion propagation on the manifold while the Kalman filter operates on a minimal local error.")
    add_equation(doc, "x = [p_n, v_n, q_np, b_a, b_g]     delta-x = [delta-p, delta-v, delta-theta, delta-b_a, delta-b_g]", "Nominal state and 15-dimensional error state.")
    add_heading(doc, "7.2 Strapdown propagation", 2)
    add_equation(doc, "q(k+1) = q(k) x Exp((omega_m - b_g) * dt)")
    add_equation(doc, "a_n = R(q(k)) * (a_m - b_a) + g_n")
    add_equation(doc, "v(k+1) = v(k) + a_n * dt")
    add_equation(doc, "p(k+1) = p(k) + v(k) * dt + 0.5 * a_n * dt^2")
    add_body(doc, "The implementation must normalize the quaternion after propagation and use a consistent gravity sign convention. Biases are modeled as random walks. The continuous error dynamics are discretized to obtain the transition matrix F and noise-input matrix G. Covariance propagation follows the standard form below.")
    add_equation(doc, "P(k+1|k) = F(k) * P(k|k) * F(k)^T + G(k) * Q(k) * G(k)^T", "Q includes accelerometer noise, gyroscope noise and bias random-walk densities.")
    add_heading(doc, "7.3 Generic measurement update", 2)
    add_equation(doc, "r = z - h(x)     S = HPH^T + R     K = PH^T S^-1")
    add_equation(doc, "delta-x = K r     P = (I - KH)P(I - KH)^T + K R K^T", "Joseph-form covariance update is preferred for numerical stability.")
    add_body(doc, "After computing the local error, the nominal position, velocity, orientation and biases are corrected. The attitude error is injected through a small rotation, and the local error state is reset. Each measurement source has its own residual, Jacobian, covariance, gating threshold and reason code.")
    add_heading(doc, "7.4 Filter tuning and consistency", 2)
    for item in [
        "Estimate sensor noise and bias instability from stationary logs and Allan-deviation analysis where feasible.",
        "Tune process noise on the training set, then freeze it before evaluating held-out trips.",
        "Track **Normalized Innovation Squared (NIS)** for each measurement type and **Normalized Estimation Error Squared (NEES)** when reference state is available.",
        "Use covariance inflation when the AI output and inertial propagation share the same IMU window, because their errors are correlated.",
        "Detect numerical failure through non-positive covariance, quaternion norm error, NaN checks and innovation explosions."
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. AI Motion Model", 1)
    add_heading(doc, "8.1 Role of learning", 2)
    add_body(doc, "The neural network is not a replacement for the navigation filter. Its role is to learn patterns that are difficult to model analytically, including phone vibration, temporal signatures of acceleration and braking, and device-specific bias behavior. The first prototype predicts **forward displacement over a short window**, a corresponding **log variance**, and a motion-state distribution. Vertical displacement is excluded initially because it is weakly observable from phone-only data and is not required for most road navigation.")
    add_heading(doc, "8.2 Baseline architecture", 2)
    add_table(doc, ["Component", "Recommended baseline", "Reason"], [
        ["Input", "6 channels: accelerometer + gyroscope", "Available across more IO-VNBD smartphone subsets"],
        ["Window", "5 seconds, 50 samples at 10 Hz", "Captures acceleration, turning and stop context"],
        ["Backbone", "Compact TCN", "Causal inference, efficient temporal receptive field"],
        ["Channels", "32, 64, 64", "Small enough for mobile baseline"],
        ["Kernel/dilation", "Kernel 3; dilations 1, 2, 4", "Expands temporal context without recurrence"],
        ["Regularization", "Dropout 0.1 to 0.2", "Reduces device and route overfitting"],
        ["Heads", "delta-s mean, log variance, motion class", "Supports uncertainty-aware fusion and constraints"],
    ], [1700, 3100, 4560], first_col_bold=True)
    add_heading(doc, "8.3 Loss and uncertainty", 2)
    add_equation(doc, "L_motion = 0.5 * (((y - mu)^2 / sigma^2) + log(sigma^2))", "Gaussian negative log likelihood encourages accurate and calibrated uncertainty.")
    add_equation(doc, "L_total = L_motion + lambda_cls * L_classification + lambda_smooth * L_smoothness")
    add_body(doc, "The variance head must be bounded for numerical stability, for example by predicting log variance and clamping it to a validated range. The ESKF converts the predicted displacement to a measurement and uses the predicted variance as part of R. If the variance is high, the update is automatically down-weighted or rejected.")
    add_heading(doc, "8.4 Motion classes", 2)
    for item in [
        "**Stationary:** supports stop detection but requires confirmation from low gyroscope norm and acceleration variance.",
        "**Cruising:** forward-motion model with normal non-holonomic constraint strength.",
        "**Accelerating and braking:** higher longitudinal dynamics; speed change is expected.",
        "**Turning:** lateral acceleration and yaw rate rise; relax overly strict lateral-velocity constraints.",
        "**Anomalous or uncertain:** phone movement, pothole, saturation or out-of-distribution window; suppress learned update."
    ]:
        add_bullet(doc, item)
    add_heading(doc, "8.5 Mobile optimization", 2)
    add_body(doc, "Training is performed in PyTorch. The frozen model is exported to ONNX with fixed tensor names, supported operators and a reproducible preprocessing contract. ONNX Runtime Mobile provides portable inference, but **runtime portability is not sensor independence**. Quantization to INT8 should be attempted only after measuring accuracy loss, calibration drift and latency improvement on the target phone.")

    add_heading(doc, "9. Dataset, Training and Generalization", 1)
    add_heading(doc, "9.1 Primary dataset", 2)
    add_body(doc, "The primary public source is **IO-VNBD**, which contains vehicle-navigation data collected with both vehicle/reference equipment and smartphones. The reported aggregate is approximately **98 hours and 5,700 km**, consisting of about **40 hours and 1,300 km of vehicle-extracted data** and **58 hours and 4,400 km of smartphone data**. These values should be confirmed against the exact repository release used in the experiment and recorded in a data manifest.")
    add_callout(doc, "Data caution", "Not every subset exposes the same sensor fields. Some smartphone recordings may omit reliable magnetometer or orientation channels. The baseline should therefore use **accelerometer plus gyroscope**, with magnetometer treated as an optional ablation.", fill=LIGHT_ORANGE, color=GOLD)
    add_heading(doc, "9.2 Dataset audit manifest", 2)
    add_table(doc, ["Field", "What must be recorded"], [
        ["Subset identity", "Repository version, route/trip ID and source folder"],
        ["Platform", "Phone model, operating system, mounting and sampling rates"],
        ["Coverage", "Country/region, duration, distance and road type"],
        ["IMU schema", "Axes, units, timestamps, missing channels and calibration status"],
        ["Reference", "Reference trajectory source, rate, expected accuracy and synchronization"],
        ["Quality", "Gaps, duplicates, saturation, abnormal timestamps and unusable intervals"],
        ["Split assignment", "Train, validation or test, with leakage check"],
    ], [2200, 7160], first_col_bold=True)
    add_heading(doc, "9.3 Label generation", 2)
    add_body(doc, "Reference positions are transformed into the same local ENU frame used by the estimator. Forward displacement labels are computed over each input window by projecting the reference displacement onto the vehicle-forward axis at the window start or by integrating reference speed with a documented convention. Label timestamps must be interpolated to the IMU clock after estimating clock offset. Samples that straddle reference gaps or uncertain alignment are excluded.")
    add_heading(doc, "9.4 Leakage-safe splitting", 2)
    add_body(doc, "Randomly shuffling overlapping windows is unacceptable because adjacent windows share most of their samples and route context. Splits must be created at the **trip, driver, phone and geography level**. A strong test plan includes a leave-one-device-out split and a country or region holdout. All normalization statistics are computed from the training split only.")
    add_heading(doc, "9.5 Augmentation", 2)
    for item in [
        "Add constant and slowly varying accelerometer and gyroscope bias.",
        "Apply per-axis scale-factor and small axis-misalignment perturbations.",
        "Add realistic timestamp jitter, short packet loss and sensor dropouts.",
        "Rotate the phone mounting orientation while preserving the declared vehicle frame.",
        "Inject vibration bursts, pothole-like impulses, saturation and magnetometer corruption.",
        "Vary sampling rate within the supported range and test the resampling pipeline."
    ]:
        add_bullet(doc, item)
    add_heading(doc, "9.6 Training protocol", 2)
    add_number(doc, "Fit preprocessing statistics and augmentation ranges on training trips only.")
    add_number(doc, "Train the TCN using early stopping on validation negative log likelihood and drift impact, not only pointwise MAE.")
    add_number(doc, "Calibrate predictive variance using validation residuals, then freeze the model and calibration parameters.")
    add_number(doc, "Export to ONNX and verify numerical parity against PyTorch on a fixed test bundle.")
    add_number(doc, "Run the complete estimator on untouched trips and simulated denied intervals.")

    add_heading(doc, "10. GNSS Integrity and Denied-Mode Operation", 1)
    add_heading(doc, "10.1 Trust inputs", 2)
    add_body(doc, "GNSS trust is evaluated using multiple indicators because no single Android accuracy value is sufficient. Inputs include reported horizontal and vertical accuracy, fix age, satellite count, carrier-to-noise density where raw measurements are accessible, speed and bearing consistency, unrealistic jumps, disagreement with inertial prediction and ESKF innovation statistics.")
    add_heading(doc, "10.2 Innovation gate", 2)
    add_equation(doc, "d^2 = r^T * S^-1 * r", "Normalized Innovation Squared. Reject the update when d^2 exceeds the selected chi-square threshold.")
    add_body(doc, "The threshold depends on the measurement dimension and chosen false-alarm probability. A single rejection does not necessarily imply spoofing, but repeated large innovations combined with quality or motion inconsistencies move the state to **SUSPECTED_SPOOF**. Thresholds must be validated on genuine urban multipath so the detector does not confuse ordinary degradation with a deliberate attack.")
    add_heading(doc, "10.3 Complete denial", 2)
    add_body(doc, "When GNSS is denied, no synthetic or stale fix is inserted into the ESKF. The filter performs IMU propagation and accepts only learned, stop, vehicle and high-confidence map measurements. The covariance is allowed to grow. The application displays the elapsed denied time, estimated uncertainty and current confidence level. If uncertainty exceeds the ODD threshold, route guidance becomes advisory or is suppressed.")
    add_heading(doc, "10.4 Safe recovery", 2)
    for item in [
        "Require multiple consecutive fixes that pass quality and innovation checks.",
        "Check consistency of position, speed and heading against the predicted state.",
        "Apply a gradual correction or covariance-aware smoothing rather than an instantaneous map jump.",
        "Preserve the pre-recovery trajectory and diagnostics for post-run evaluation.",
        "Return to TRUSTED only after a configurable dwell time and stable innovations."
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Vehicle Constraints and Map Matching", 1)
    add_heading(doc, "11.1 Non-holonomic constraint", 2)
    add_body(doc, "A normal road vehicle has near-zero lateral and vertical velocity in its own frame during ordinary driving. Percorsa uses this as a pseudo-measurement, not as a permanent equality. The measurement covariance is relaxed during sharp turns, skids, rough roads, banked roads, off-road motion or suspected phone movement.")
    add_equation(doc, "z_NHC = [v_y in F_v, v_z in F_v] approximately [0, 0]", "Constraint strength is controlled by a validity probability and adaptive covariance.")
    add_heading(doc, "11.2 Stop detection and zero-velocity update", 2)
    add_body(doc, "A stop is declared only when several signals agree over a minimum duration: low accelerometer variance after gravity handling, low gyroscope norm, low predicted motion probability and low estimated speed. Engine vibration and a phone resting loosely in a holder can violate simple thresholds, so the detector must be evaluated while idling. A valid stop update constrains velocity toward zero and improves bias observability.")
    add_heading(doc, "11.3 HMM/Viterbi map matcher", 2)
    add_body(doc, "The road matcher keeps multiple candidate road segments for each estimated position. The **emission score** uses perpendicular distance, heading difference, road class and the ESKF covariance. The **transition score** uses graph connectivity, path distance, turn feasibility, direction restrictions and consistency with traveled distance. Viterbi decoding selects the most probable path over a short history, while top-k hypotheses preserve ambiguity at parallel roads or intersections.")
    add_table(doc, ["Map signal", "Used for", "Failure protection"], [
        ["Perpendicular distance", "Candidate emission likelihood", "Scale by position covariance"],
        ["Heading difference", "Reject opposing or crossing roads", "Relax at low speed and junctions"],
        ["Connectivity", "Plausible path sequence", "Retain top-k candidates"],
        ["Road direction", "One-way feasibility", "Disable when map attribute missing"],
        ["Travel distance", "Transition consistency", "Use learned/INS uncertainty"],
        ["Match confidence", "Authorize pseudo-measurement", "No update below threshold"],
    ], [2100, 3250, 4010], first_col_bold=True)
    add_callout(doc, "Safety rule", "A map is contextual evidence, not ground truth. **Never force the state onto a road when the match is ambiguous.** Low-confidence matching should remain a display hypothesis and should not correct the ESKF.", fill=LIGHT_ORANGE, color=RED)

    add_heading(doc, "12. Android and Edge Implementation", 1)
    add_heading(doc, "12.1 Runtime architecture", 2)
    for item in [
        "An Android foreground service owns sensor subscriptions and continues acquisition with the screen off.",
        "A lock-free or bounded queue carries timestamped samples to preprocessing. Queue overflow is counted and exposed.",
        "The estimator thread runs at the fixed 10 Hz profile and processes all samples up to the current fusion timestamp.",
        "The ONNX model executes on a worker thread after one warm-up pass. Preallocated tensors reduce garbage collection.",
        "Offline OpenStreetMap tiles and a compact road graph are stored per region.",
        "State snapshots are persisted safely so a process restart is visible and does not masquerade as continuous navigation.",
        "The UI renders mode, route, covariance/confidence and diagnostics without blocking estimation."
    ]:
        add_bullet(doc, item)
    add_heading(doc, "12.2 Timing budget", 2)
    add_table(doc, ["Stage", "Prototype budget at 10 Hz", "Measurement"], [
        ["Sensor ingest/resampling", "Less than 10 ms", "p50/p95 latency and sample gaps"],
        ["TCN inference", "Less than 25 ms", "warm and cold latency on target phone"],
        ["INS + ESKF", "Less than 10 ms", "per-cycle computation time"],
        ["Map matching", "Less than 30 ms average", "candidate count and p95 latency"],
        ["UI/logging reserve", "Less than 15 ms", "frame drops and queue depth"],
        ["Total", "Less than 100 ms", "sustained 10 Hz under thermal load"],
    ], [2500, 2650, 4210], first_col_bold=True)
    add_heading(doc, "12.3 Hardware profiles", 2)
    add_body(doc, "The report separates two profiles to remove the earlier smartphone and FOG contradiction. **Profile A is the SIH prototype:** Android smartphone sensors, 10 Hz fusion and no mandatory external hardware. **Profile B is an optional scale-up path:** an external MEMS or FOG IMU sampled up to 200 Hz on a companion edge computer. Profile B requires separate calibration, retraining, noise parameters and validation and is not evidence for Profile A performance.")
    add_table(doc, ["Profile", "Sensors", "Compute", "Claim"], [
        ["A: Smartphone", "Phone accelerometer, gyroscope, optional gated magnetometer and GNSS", "Android + ONNX Runtime Mobile", "Primary demonstrable prototype at 10 Hz"],
        ["B: External IMU", "Higher-rate MEMS or FOG plus optional GNSS", "Companion edge device", "Future extension, validated separately up to 200 Hz"],
    ], [1700, 3900, 2300, 1460], first_col_bold=True, font_size=8.8)

    add_heading(doc, "13. Experimental Methodology", 1)
    add_heading(doc, "13.1 Baselines and ablations", 2)
    add_table(doc, ["ID", "Configuration", "Purpose"], [
        ["B0", "Last trusted fix held constant", "Minimum reference behavior"],
        ["B1", "Constant velocity and heading", "Simple navigation extrapolation"],
        ["B2", "Pure smartphone INS", "Quantify unaided inertial drift"],
        ["B3", "INS + NHC/ZUPT", "Measure physics-constraint benefit"],
        ["B4", "AI-only displacement integration", "Measure learned odometry alone"],
        ["B5", "AI + INS + ESKF", "Measure hybrid estimator"],
        ["B6", "Full system + map", "Measure final confidence-gated system"],
        ["A1", "Full system without uncertainty head", "Test uncertainty contribution"],
        ["A2", "Full system without phone-frame adaptation", "Test mounting robustness"],
        ["A3", "Full system with optional magnetometer", "Measure benefit versus disturbance risk"],
    ], [800, 4100, 4460], first_col_bold=True)
    add_heading(doc, "13.2 Denial scenarios", 2)
    add_table(doc, ["Dimension", "Test values"], [
        ["Time", "5, 10, 30, 60 and 120 seconds"],
        ["Distance", "50, 100, 500 and 1,000 meters"],
        ["Road geometry", "Straight, gradual turn, sharp turn, intersection, parallel road and loop"],
        ["Motion", "Cruise, stop-go, braking, acceleration, idle and rough-road vibration"],
        ["Signal", "Complete denial, urban degradation, delayed fixes and spoof-like jumps"],
        ["Generalization", "Unseen trip, phone, driver, route and geography"],
        ["Mounting", "Portrait, landscape and changed mount after calibration"],
    ], [2100, 7260], first_col_bold=True)
    add_heading(doc, "13.3 Accuracy metrics", 2)
    add_equation(doc, "Endpoint drift (%) = 100 * ||p_est(end) - p_ref(end)|| / distance_travelled")
    add_equation(doc, "ATE_RMSE = sqrt((1/N) * sum(||p_est(i) - p_ref(i)||^2))")
    add_equation(doc, "Heading error = wrap_to_pi(heading_est - heading_ref)")
    add_body(doc, "Report mean, median, p90, p95 and worst case. Distance-normalized drift can be unstable for very short or nearly stationary segments, so always report the absolute endpoint error and traveled distance beside the percentage.")
    add_heading(doc, "13.4 Uncertainty and integrity metrics", 2)
    for item in [
        "Negative log likelihood of position or displacement residuals.",
        "Coverage of 50%, 90% and 95% confidence regions.",
        "Median confidence-ellipse area and its growth during denial.",
        "NIS acceptance rate for trusted GNSS and rejection rate for injected bad fixes.",
        "False rejection rate under genuine multipath and false acceptance rate under spoof-like jumps.",
        "Time from GNSS return to safe TRUSTED recovery."
    ]:
        add_bullet(doc, item)
    add_heading(doc, "13.5 Mobile metrics", 2)
    add_body(doc, "Measure model size, APK increase, p50 and p95 inference latency, full-loop latency, achieved update rate, dropped samples, peak RAM, average CPU load, battery drain per hour and thermal throttling over at least a 30-minute replay or drive. Report phone model, OS version, power mode and ambient conditions.")
    add_heading(doc, "13.6 Statistical protocol", 2)
    add_body(doc, "Each method must be evaluated on identical denial intervals. Use paired comparisons per interval and bootstrap confidence intervals across trips, not across overlapping windows. Hyperparameters and thresholds are frozen before the final test. Failure cases are retained in the aggregate rather than removed as outliers unless a pre-declared data-quality rule excludes them.")

    add_heading(doc, "14. Results and Ablation Reporting", 1)
    add_callout(doc, "Evidence status", "The current concept report does **not contain independently measured final performance values**. The tables in this section define the exact evidence that must be inserted after implementation. This prevents design targets from being misrepresented as achieved results.", fill=LIGHT_ORANGE, color=RED)
    add_heading(doc, "14.1 Main results table", 2)
    add_table(doc, ["Method", "Median drift %", "p95 drift %", "Heading p95", "95% coverage", "Status"], [
        ["B0 Last fix", "Not measured", "Not measured", "Not measured", "N/A", "Run evaluation"],
        ["B2 Pure INS", "Not measured", "Not measured", "Not measured", "Not measured", "Run evaluation"],
        ["B3 INS + constraints", "Not measured", "Not measured", "Not measured", "Not measured", "Run evaluation"],
        ["B5 Hybrid ESKF", "Not measured", "Not measured", "Not measured", "Not measured", "Run evaluation"],
        ["B6 Full system", "Not measured", "Not measured", "Not measured", "Not measured", "Run evaluation"],
    ], [2100, 1450, 1450, 1450, 1450, 1460], first_col_bold=True, font_size=8.2)
    add_heading(doc, "14.2 Required figures", 2)
    for item in [
        "Reference versus estimated trajectory for one straight, one turning and one parallel-road case.",
        "Error versus denied time with median and p90 bands.",
        "Covariance growth and true error on the same time axis to expose overconfidence.",
        "GNSS innovation and Trust Manager state during normal degradation and injected jumps.",
        "Ablation chart showing how NHC, AI update and map matching change endpoint drift.",
        "Mobile latency histogram and 30-minute thermal/battery trace."
    ]:
        add_bullet(doc, item)
    add_heading(doc, "14.3 Acceptance decision", 2)
    add_body(doc, "The less-than-10% drift target should be declared achieved only if the pre-registered primary test split meets the threshold at the selected aggregate, preferably median and p90, and if uncertainty coverage is not severely overconfident. If the target is met only on short intervals or familiar devices, the claim must name that condition explicitly.")

    add_heading(doc, "15. Feasibility, Risks and Limitations", 1)
    add_heading(doc, "15.1 Feasibility assessment", 2)
    add_table(doc, ["Area", "Technical basis", "Bounded claim"], [
        ["Hardware", "Existing phone accelerometer and gyroscope; optional trusted GNSS", "No mandatory OBD/CAN for primary profile"],
        ["Data", "IO-VNBD provides multi-hour, multi-route smartphone and reference logs", "Exact usable coverage follows dataset audit"],
        ["Algorithm", "INS, ESKF, TCN, NHC and HMM map matching are established components", "Innovation is their integrity-aware integration"],
        ["Performance", "10 Hz model and filter are plausible on modern phones", "Must be demonstrated on declared target devices"],
        ["Prototype", "Android logging, offline replay and ONNX inference are implementable", "Live road test follows replay validation"],
        ["Deployment", "Modular services and offline maps support field use", "External FOG is a separate future profile"],
    ], [1600, 4500, 3260], first_col_bold=True, font_size=8.7)
    add_heading(doc, "15.2 Risk register", 2)
    add_table(doc, ["Risk", "Detection", "Mitigation", "Residual limitation"], [
        ["Phone moved after alignment", "Frame-change classifier and abrupt gravity/gyro inconsistency", "Freeze corrections; request re-alignment", "Trajectory may be unavailable until alignment restored"],
        ["Long complete denial", "Covariance and drift bound exceed limit", "Mark UNSAFE; request new anchor", "Phone-only drift remains unbounded over time"],
        ["Magnetic interference", "Field-norm and innovation tests", "Reject magnetometer", "Yaw relies on gyro, motion and road context"],
        ["Parallel-road ambiguity", "Low HMM margin between top candidates", "Retain top-k and avoid map update", "Absolute road identity may remain unknown"],
        ["Bad or spoofed GNSS", "Quality, motion and NIS inconsistency", "Reject and enter suspected-spoof state", "Sophisticated slow spoofing is harder to detect"],
        ["Engine vibration/potholes", "High-frequency energy and anomaly score", "Adaptive noise and update suppression", "Some dynamic events resemble motion changes"],
        ["Unseen phone model", "Out-of-distribution and calibration checks", "Device holdout training and covariance inflation", "Revalidation is required"],
        ["Map errors", "Topology conflict and low match confidence", "Disable pseudo-measurement", "No map-aided correction in affected area"],
    ], [1950, 2350, 2440, 2620], first_col_bold=True, font_size=7.9)
    add_heading(doc, "15.3 Known limitations", 2)
    add_body(doc, "Percorsa does not guarantee lane-level accuracy, does not infer an absolute origin without an anchor, and does not eliminate inertial drift. Two-wheelers have materially different roll and lateral dynamics and should not be included in the primary claim without separate modeling. Performance learned from IO-VNBD may not transfer to every phone, road surface or driving style. Map matching can improve a route estimate but can also introduce error if applied with excessive confidence.")

    add_heading(doc, "16. Privacy, Security and Deployment", 1)
    add_heading(doc, "16.1 Privacy", 2)
    for item in [
        "Process sensor and trajectory data **on device by default**.",
        "Require explicit consent before exporting trip logs for debugging or retraining.",
        "Remove direct identifiers and coarsen or trim home/work endpoints before dataset contribution.",
        "Encrypt local logs and apply a retention period that the user can control.",
        "Document every collected field, purpose and retention rule in the application."
    ]:
        add_bullet(doc, item)
    add_heading(doc, "16.2 Security", 2)
    for item in [
        "Sign application and model updates and verify their integrity before activation.",
        "Protect offline maps and calibration files from unauthorized modification.",
        "Treat GNSS and map inputs as untrusted external data and validate ranges, timestamps and identifiers.",
        "Rate-limit diagnostic export and avoid logging secrets or stable device identifiers.",
        "Maintain an audit trace of mode changes and rejected measurements for forensic review."
    ]:
        add_bullet(doc, item)
    add_heading(doc, "16.3 User interface safety", 2)
    add_body(doc, "The map interface must separate **estimated route** from **trusted absolute position**. Confidence should be communicated with a clear label and an uncertainty region, not a misleading fixed-size blue dot. During UNSAFE mode, the user should see that precision is unavailable and should not receive turn-by-turn instructions that depend on an uncertain road identity.")

    add_heading(doc, "17. Implementation Roadmap", 1)
    add_table(doc, ["Phase", "Engineering deliverable", "Exit evidence"], [
        ["1. Data foundation", "IO-VNBD audit, synchronized local-frame sequences, leakage-safe splits", "Manifest, quality report and reproducible preprocessing"],
        ["2. Classical baseline", "Strapdown INS, 15-state ESKF, replay tool", "Pure INS and GNSS/INS baseline metrics"],
        ["3. Learned motion", "TCN displacement, variance and motion state", "Held-out accuracy, calibration and ONNX parity"],
        ["4. Constraints", "NHC, stop detector and confidence-gated HMM matcher", "Ablation results and failure traces"],
        ["5. Integrity", "GNSS Trust Manager and recovery state machine", "Injected degradation/spoof tests"],
        ["6. Android", "Foreground service, ONNX inference, offline map UI", "10 Hz sustained device demonstration"],
        ["7. Field validation", "Controlled tunnel/parking/urban routes", "Signed test report with all primary metrics"],
        ["8. SIH demo", "Live denied-mode journey and replayable failure cases", "Reproducible demo script and fallback recording"],
    ], [1500, 4780, 3080], first_col_bold=True, font_size=8.6)
    add_heading(doc, "17.1 Minimum viable demonstration", 2)
    add_number(doc, "Initialize outdoors with trusted GNSS and a fixed phone mount.")
    add_number(doc, "Enter a controlled denial segment or switch to a pre-recorded synchronized replay.")
    add_number(doc, "Show continuous estimated position, mode, heading and confidence while GNSS input is absent.")
    add_number(doc, "Display pure INS and Percorsa trajectories together to show the effect of learning and constraints.")
    add_number(doc, "Restore GNSS and demonstrate innovation-gated, smooth recovery.")
    add_number(doc, "Open the diagnostic view to show accepted/rejected updates and measured device latency.")
    add_heading(doc, "17.2 SIH evidence package", 2)
    for item in [
        "Versioned source code and frozen model checksum.",
        "Dataset manifest and exact train/validation/test trip IDs.",
        "Configuration file containing noise, gates and map thresholds.",
        "Automated replay script and a small redistributable test bundle.",
        "Results tables, trajectory plots, failure cases and mobile profiling report.",
        "One-page claim sheet separating **measured result**, **design target** and **future work**."
    ]:
        add_bullet(doc, item)

    add_heading(doc, "18. Conclusion", 1)
    add_body(doc, "Percorsa is technically feasible as a **bounded, smartphone-first dead-reckoning prototype** if the team treats uncertainty, data leakage, mounting variation and map ambiguity as first-class engineering problems. Its strongest differentiator is not the use of a neural network alone. It is the end-to-end integrity design that decides when to trust GNSS, when to rely on inertial and learned motion, when to apply road constraints and when to admit that the estimate is no longer safe.")
    add_body(doc, "The next milestone is evidence. The team should first establish a reproducible classical baseline, then add the TCN, constraints and trust manager one at a time. Every improvement must be supported by held-out trajectory metrics, calibrated uncertainty and measured mobile performance. This produces a credible SIH submission without overclaiming what phone-only inertial navigation can achieve.")

    page_break(doc)
    add_heading(doc, "References", 1)
    refs = [
        "IO-VNBD repository and dataset documentation. https://github.com/onyekpeu/IO-VNBD",
        "Android Developers, Sensors Overview and SensorEvent timestamp documentation. https://developer.android.com/develop/sensors-and-location/sensors/sensors_overview",
        "ONNX Runtime, Mobile deployment documentation. https://onnxruntime.ai/docs/get-started/with-mobile.html",
        "OpenStreetMap project and data documentation. https://www.openstreetmap.org/ and https://wiki.openstreetmap.org/",
        "Barfoot, T. D. State Estimation for Robotics. Cambridge University Press, 2017.",
        "Groves, P. D. Principles of GNSS, Inertial, and Multisensor Integrated Navigation Systems. Artech House, 2nd edition, 2013.",
        "Maybeck, P. S. Stochastic Models, Estimation, and Control, Volume 1. Academic Press, 1979.",
        "Original SIH26168 problem statement and Team Percorsa concept material supplied for this report."
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(f"[{i}] ")
        set_run_font(r, size=9.5, bold=True, color=NAVY)
        r = p.add_run(ref)
        set_run_font(r, size=9.5, color=INK)

    add_heading(doc, "Appendix A. Interface Contracts", 1)
    add_table(doc, ["Message", "Minimum fields"], [
        ["ImuSample", "timestamp_ns, ax, ay, az, gx, gy, gz, quality_flags"],
        ["GnssObservation", "timestamp_ns, lat, lon, alt, speed, bearing, accuracies, satellite metrics"],
        ["MotionPrediction", "window_end_ns, delta_s_mean, delta_s_variance, state_probabilities, OOD score"],
        ["NavigationState", "timestamp_ns, ENU position, velocity, quaternion, biases, covariance, mode"],
        ["MapHypothesis", "road_edge_id, probability, along_track, heading_residual, candidate_rank"],
        ["DiagnosticEvent", "timestamp_ns, source, accepted, reason_code, residual, NIS, threshold"],
    ], [2400, 6960], first_col_bold=True)
    add_heading(doc, "Appendix B. Configuration Parameters", 1)
    add_table(doc, ["Group", "Parameters to version"], [
        ["Sensor", "target rate, filter cutoffs, saturation limits, gap limits"],
        ["ESKF", "initial covariance, process noise, bias random walk, gravity constant"],
        ["AI", "window length, normalization, model checksum, variance clamp"],
        ["GNSS", "accuracy gates, NIS threshold, consecutive recovery count, dwell time"],
        ["Constraints", "NHC covariance, stop thresholds, validity probabilities"],
        ["Map", "search radius, emission weights, transition weights, top-k, confidence gate"],
        ["Safety", "maximum uncertainty, maximum denied time for each application profile"],
    ], [2200, 7160], first_col_bold=True)
    add_heading(doc, "Appendix C. Verification Checklist", 1)
    checks = [
        "All coordinate frames, units and quaternion conventions are declared.",
        "No GNSS update enters the ESKF without a trust decision and reason code.",
        "Phone movement and timestamp gaps produce a visible estimator response.",
        "Train, validation and test trips have no overlapping windows or shared route leakage beyond the declared protocol.",
        "PyTorch and ONNX outputs match within the selected numerical tolerance.",
        "Every target claim is labeled as measured, planned or future work.",
        "All reported metrics identify the dataset split, denial interval and aggregation method.",
        "The Android test records latency, dropped samples, CPU, memory, battery and temperature.",
        "Low-confidence road matching cannot force an ESKF correction.",
        "An uncertainty threshold drives the UNSAFE state and user-interface warning."
    ]
    for c in checks:
        add_bullet(doc, "[ ] " + c)

    # Document metadata and compatibility.
    props = doc.core_properties
    props.title = "Percorsa - SIH26168 Detailed Technical Report"
    props.subject = "AI-ML Intelligent Dead Reckoning for GNSS-Denied Ground-Vehicle Navigation"
    props.author = "Team Percorsa"
    props.keywords = "SIH26168, dead reckoning, GNSS denied, smartphone IMU, ESKF, TCN, map matching"

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
